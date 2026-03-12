"""
Smart Text Converter module (V4 — Universal).
Produces high-quality DOCX from any PDF using PyMuPDF:
- Preserves original fonts, sizes, colors (bold, italic)
- Detects headings by font size
- Extracts and places images at correct positions
- Handles multi-column layouts via table-based rendering  
- Reconstructs bullet/numbered lists
- Separates header/footer from body content
- Handles page breaks between pages
"""

import logging
import os
import re
import tempfile

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

_CONTROL_CHARS = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]')
_BULLET_PATTERN = re.compile(r'^[\s]*[•●○◦▪▸►–—\-\*]\s+')
_NUMBER_PATTERN = re.compile(r'^[\s]*\d{1,3}[\.\)]\s+')
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class SmartConverter:
    """
    Universal PDF-to-DOCX converter using PyMuPDF.
    Automatically adapts to any PDF layout without hardcoded styles.
    """

    HEADER_Y_THRESHOLD = 0.06   # Top 6% = header zone
    FOOTER_Y_THRESHOLD = 0.92   # Bottom 8% = footer zone
    COLUMN_GAP_MIN = 30.0       # Min gap between columns (pts)
    MIN_IMAGE_SIZE = 15         # Min image dimension to extract (pts)

    def __init__(self):
        pass

    def convert(self, input_pdf: str, output_docx: str,
                pages: list[int] | None = None) -> str:
        logger.info("Smart V4 converting: %s -> %s", input_pdf, output_docx)

        pdf_doc = fitz.open(input_pdf)
        doc = Document()

        # Set default style
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)

        # Analyze the entire document to determine global font metrics
        font_stats = self._analyze_fonts(pdf_doc, pages)

        page_numbers = pages if pages else list(range(len(pdf_doc)))

        with tempfile.TemporaryDirectory() as tmp_dir:
            for page_idx, page_num in enumerate(page_numbers):
                if page_num >= len(pdf_doc):
                    continue

                page = pdf_doc[page_num]
                logger.info("Processing page %d/%d", page_idx + 1, len(page_numbers))

                if page_idx > 0:
                    doc.add_page_break()

                self._process_page(doc, page, pdf_doc, tmp_dir, page_num, font_stats)

        pdf_doc.close()

        # Set margins
        for section in doc.sections:
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)

        doc.save(output_docx)
        logger.info("Smart V4 conversion complete: %s", output_docx)
        return output_docx

    # ── Font Analysis ────────────────────────────────────────────

    def _analyze_fonts(self, pdf_doc: fitz.Document,
                       pages: list[int] | None) -> dict:
        """Analyze fonts across the PDF to determine heading thresholds dynamically."""
        size_counts = {}
        page_numbers = pages if pages else list(range(min(len(pdf_doc), 20)))  # Sample max 20 pages

        for page_num in page_numbers:
            if page_num >= len(pdf_doc):
                continue
            page = pdf_doc[page_num]
            text_dict = page.get_text("dict")
            for block in text_dict.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if not text:
                            continue
                        size = round(span.get("size", 10), 1)
                        size_counts[size] = size_counts.get(size, 0) + len(text)

        if not size_counts:
            return {"body_size": 10.0, "heading1_min": 16.0, "heading2_min": 12.0}

        # Body size = most common font size (by character count)
        body_size = max(size_counts, key=size_counts.get)

        # Get all sizes sorted
        sizes = sorted(size_counts.keys())
        larger_sizes = [s for s in sizes if s > body_size + 1.0]

        heading1_min = larger_sizes[-1] if len(larger_sizes) >= 1 else body_size + 6.0
        heading2_min = larger_sizes[0] if len(larger_sizes) >= 1 else body_size + 2.0
        if len(larger_sizes) >= 2:
            heading1_min = larger_sizes[-1] if larger_sizes[-1] > heading2_min + 2 else heading2_min + 2
            heading2_min = larger_sizes[0]

        logger.info("Font stats: body=%.1f, h2_min=%.1f, h1_min=%.1f",
                     body_size, heading2_min, heading1_min)

        return {
            "body_size": body_size,
            "heading1_min": heading1_min,
            "heading2_min": heading2_min,
        }

    # ── Page Processing ──────────────────────────────────────────

    def _process_page(self, doc: Document, page: fitz.Page,
                      pdf_doc: fitz.Document, tmp_dir: str,
                      page_num: int, font_stats: dict):
        """Process a single page: extract tables, images, text blocks in reading order."""
        page_h = page.rect.height
        page_w = page.rect.width

        # 1. Extract text blocks
        text_dict = page.get_text("dict")
        raw_blocks = [b for b in text_dict.get("blocks", [])
                      if b["type"] == 0 and self._block_has_text(b)]

        # 2. Separate header/body/footer
        header_y = page_h * self.HEADER_Y_THRESHOLD
        footer_y = page_h * self.FOOTER_Y_THRESHOLD

        body_blocks = [b for b in raw_blocks
                       if b["bbox"][1] >= header_y and b["bbox"][1] < footer_y]

        # 3. Detect tables using PyMuPDF find_tables()
        tables_info = self._detect_tables(page)

        # 4. Filter out text blocks that overlap with detected tables
        non_table_blocks = []
        for block in body_blocks:
            if not any(self._block_overlaps_table(block, t_info["bbox"]) for t_info in tables_info):
                non_table_blocks.append(block)

        # 5. Extract vector drawings for cell background colors
        drawings = page.get_drawings()

        # 6. Build a unified list of renderable elements sorted by Y position
        elements = []  # (y_pos, type, data)

        # Add images
        img_elements = self._collect_image_elements(page, pdf_doc, tmp_dir,
                                                     page_num, header_y, footer_y)
        elements.extend(img_elements)

        # Add tables
        for t_info in tables_info:
            y_pos = t_info["bbox"][1]  # top of table
            elements.append((y_pos, "table", t_info, drawings))

        # Add text blocks (possibly grouped by column)
        columns = self._detect_columns(non_table_blocks, page_w)
        if columns == 2:
            # Group into full-width, left, right
            mid = page_w / 2.0
            full_width = []
            left_col = []
            right_col = []
            for block in non_table_blocks:
                x0, _, x1, _ = block["bbox"]
                block_width = x1 - x0
                if block_width > page_w * 0.55:
                    full_width.append(block)
                elif x1 < mid + 20:
                    left_col.append(block)
                elif x0 > mid - 20:
                    right_col.append(block)
                else:
                    full_width.append(block)
            # Add full-width blocks individually
            for b in full_width:
                elements.append((b["bbox"][1], "text_block", b))
            # Add column pair as one element
            if left_col or right_col:
                min_y = min(
                    (left_col[0]["bbox"][1] if left_col else 9999),
                    (right_col[0]["bbox"][1] if right_col else 9999)
                )
                elements.append((min_y, "columns", left_col, right_col))
        else:
            for block in non_table_blocks:
                elements.append((block["bbox"][1], "text_block", block))

        # 7. Sort by Y position and render
        elements.sort(key=lambda x: x[0])

        for elem in elements:
            if elem[1] == "table":
                self._render_table(doc, elem[2], elem[3], font_stats)
            elif elem[1] == "text_block":
                self._add_block(doc, elem[2], font_stats)
            elif elem[1] == "columns":
                left_col = sorted(elem[2], key=lambda b: b["bbox"][1])
                right_col = sorted(elem[3], key=lambda b: b["bbox"][1])
                self._add_column_table(doc, left_col, right_col, font_stats)
            elif elem[1] == "image":
                self._add_image_element(doc, elem[2])

    # ── Table Detection ──────────────────────────────────────────

    def _detect_tables(self, page: fitz.Page) -> list[dict]:
        """Detect tables using PyMuPDF's find_tables()."""
        try:
            tables_result = page.find_tables()
            tables_info = []
            for table in tables_result.tables:
                extracted = table.extract()
                if not extracted or len(extracted) < 2:
                    continue  # Skip trivial tables
                tables_info.append({
                    "bbox": table.bbox,  # (x0, y0, x1, y1)
                    "rows": extracted,
                    "row_count": len(extracted),
                    "col_count": table.col_count,
                })
            logger.info("Detected %d table(s) on page", len(tables_info))
            return tables_info
        except Exception as e:
            logger.warning("Table detection failed: %s", str(e)[:80])
            return []

    def _block_overlaps_table(self, block: dict, table_bbox: tuple) -> bool:
        """Check if a text block significantly overlaps with a table bounding box."""
        bx0, by0, bx1, by1 = block["bbox"]
        tx0, ty0, tx1, ty1 = table_bbox

        # Check overlap
        ox0 = max(bx0, tx0)
        oy0 = max(by0, ty0)
        ox1 = min(bx1, tx1)
        oy1 = min(by1, ty1)

        if ox0 >= ox1 or oy0 >= oy1:
            return False  # No overlap

        block_area = max((bx1 - bx0) * (by1 - by0), 1)
        overlap_area = (ox1 - ox0) * (oy1 - oy0)

        return (overlap_area / block_area) > 0.5

    # ── Table Rendering ──────────────────────────────────────────

    def _render_table(self, doc: Document, table_info: dict,
                      drawings: list, font_stats: dict):
        """Render a detected table as a real Word table with colors."""
        rows_data = table_info["rows"]
        row_count = table_info["row_count"]
        col_count = table_info["col_count"]
        table_bbox = table_info["bbox"]

        logger.info("Rendering table: %d rows x %d cols", row_count, col_count)

        table = doc.add_table(rows=row_count, cols=col_count)
        table.autofit = True

        # Get cell background colors from vector drawings
        cell_colors = self._get_cell_colors(table_bbox, row_count, col_count,
                                             rows_data, drawings)

        # Set nice table borders
        self._set_table_borders(table, show=True)

        for row_idx, row_data in enumerate(rows_data):
            for col_idx in range(min(len(row_data), col_count)):
                cell = table.cell(row_idx, col_idx)
                cell_text = row_data[col_idx] or ""
                cell_text = self._sanitize(cell_text).strip()

                # Set cell text
                para = cell.paragraphs[0]
                if cell_text:
                    run = para.add_run(cell_text)
                    run.font.size = Pt(font_stats["body_size"])
                    run.font.name = "Calibri"
                    # Bold for header row
                    if row_idx == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    # Center bullet marks
                    if cell_text.strip() in ("●", "○", "◦", "•"):
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Apply background color
                color_key = (row_idx, col_idx)
                if color_key in cell_colors:
                    self._set_cell_shading(cell, cell_colors[color_key])
                elif row_idx == 0:
                    # Default header color if no exact match
                    self._set_cell_shading(cell, "008C99")

    def _get_cell_colors(self, table_bbox: tuple, row_count: int,
                         col_count: int, rows_data: list,
                         drawings: list) -> dict:
        """Extract cell background colors from vector drawings (filled rectangles)."""
        colors = {}  # (row_idx, col_idx) -> hex_color

        tx0, ty0, tx1, ty1 = table_bbox
        col_width = (tx1 - tx0) / max(col_count, 1)
        row_height = (ty1 - ty0) / max(row_count, 1)

        # Collect filled rectangles that overlap the table
        filled_rects = []
        for d in drawings:
            fill = d.get("fill")
            if fill is None:
                continue
            rect = d.get("rect")
            if rect is None:
                continue
            # Check if drawing overlaps table area
            dr = fitz.Rect(rect)
            tr = fitz.Rect(tx0, ty0, tx1, ty1)
            if dr.intersects(tr) and dr.width > 10 and dr.height > 10:
                filled_rects.append((dr, fill))

        # For each cell, find the best matching filled rectangle
        for row_idx in range(row_count):
            for col_idx in range(col_count):
                cell_x0 = tx0 + col_idx * col_width
                cell_y0 = ty0 + row_idx * row_height
                cell_x1 = cell_x0 + col_width
                cell_y1 = cell_y0 + row_height
                cell_center_x = (cell_x0 + cell_x1) / 2
                cell_center_y = (cell_y0 + cell_y1) / 2

                for dr, fill in filled_rects:
                    if dr.contains(fitz.Point(cell_center_x, cell_center_y)):
                        hex_color = self._rgb_tuple_to_hex(fill)
                        # Skip pure white backgrounds
                        if hex_color != "FFFFFF":
                            colors[(row_idx, col_idx)] = hex_color
                        break

        return colors

    @staticmethod
    def _rgb_tuple_to_hex(rgb_tuple) -> str:
        """Convert an RGB float tuple (0-1) to a hex string."""
        r = int(min(max(rgb_tuple[0], 0), 1) * 255)
        g = int(min(max(rgb_tuple[1], 0), 1) * 255)
        b = int(min(max(rgb_tuple[2], 0), 1) * 255)
        return f"{r:02X}{g:02X}{b:02X}"

    @staticmethod
    def _set_cell_shading(cell, hex_color: str):
        """Apply background shading to a Word table cell."""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), hex_color)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    # ── Image Collection (for interleaved rendering) ─────────────

    def _collect_image_elements(self, page: fitz.Page, pdf_doc: fitz.Document,
                                 tmp_dir: str, page_num: int,
                                 header_y: float, footer_y: float) -> list:
        """Collect images as renderable elements with Y positions."""
        img_info_list = page.get_image_info(xrefs=True)
        seen_xrefs = set()
        elements = []

        for info in img_info_list:
            xref = info.get("xref", 0)
            if xref in seen_xrefs or xref == 0:
                continue
            seen_xrefs.add(xref)

            bbox = info.get("bbox", [])
            if not bbox or len(bbox) < 4:
                continue

            w_pt = abs(bbox[2] - bbox[0])
            h_pt = abs(bbox[3] - bbox[1])
            if w_pt < self.MIN_IMAGE_SIZE or h_pt < self.MIN_IMAGE_SIZE:
                continue

            img_y = bbox[1]
            if img_y < header_y or img_y >= footer_y:
                continue

            img_path = os.path.join(tmp_dir, f"img_p{page_num}_x{xref}.png")
            try:
                clip = fitz.Rect(bbox) & page.rect
                if clip.is_empty or clip.width < 5 or clip.height < 5:
                    continue
                scale = min(3.0, max(2.0, 600.0 / max(clip.width, clip.height)))
                mat = fitz.Matrix(scale, scale)
                pix = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
                pix.save(img_path)
            except Exception as e:
                logger.warning("Image xref=%d render failed: %s", xref, str(e)[:80])
                continue

            if not os.path.isfile(img_path):
                continue

            max_w_in = (page.rect.width - 80) / 72.0
            width_in = min(w_pt / 72.0, max_w_in)
            if width_in < 0.3:
                width_in = min(2.0, max_w_in)

            elements.append((img_y, "image", {
                "path": img_path,
                "width_in": width_in,
                "xref": xref,
            }))

        return elements

    def _add_image_element(self, doc: Document, img_data: dict):
        """Add an image element to the document."""
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(img_data["path"], width=Inches(img_data["width_in"]))
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)
            logger.info("Added image xref=%d (%.1f in wide)", img_data["xref"], img_data["width_in"])
        except Exception as e:
            logger.warning("Image insert failed: %s", str(e)[:80])

    # ── Column Detection ─────────────────────────────────────────

    def _detect_columns(self, blocks: list, page_width: float) -> int:
        """Detect if the page has 1 or 2 columns."""
        if not blocks:
            return 1

        mid = page_width / 2.0
        left_only = 0
        right_only = 0
        full_width = 0

        for block in blocks:
            x0, _, x1, _ = block["bbox"]
            block_width = x1 - x0

            if block_width > page_width * 0.55:
                full_width += 1
            elif x1 < mid + 20:
                left_only += 1
            elif x0 > mid - 20:
                right_only += 1
            else:
                full_width += 1

        total = left_only + right_only + full_width
        if total == 0:
            return 1

        # Need substantial content on both sides for 2-column
        if left_only >= 3 and right_only >= 3:
            return 2

        return 1

    # ── Single Column Rendering ──────────────────────────────────

    def _render_single_column(self, doc: Document, blocks: list, font_stats: dict):
        """Render blocks in single-column reading order."""
        sorted_blocks = sorted(blocks, key=lambda b: b["bbox"][1])

        for block in sorted_blocks:
            self._add_block(doc, block, font_stats)

    # ── Two Column Rendering ─────────────────────────────────────

    def _render_two_columns(self, doc: Document, blocks: list,
                            page_width: float, font_stats: dict):
        """Render two-column layout using a Word table."""
        mid = page_width / 2.0

        full_width = []
        left_col = []
        right_col = []

        for block in blocks:
            x0, _, x1, _ = block["bbox"]
            block_width = x1 - x0

            if block_width > page_width * 0.55:
                full_width.append(block)
            elif x1 < mid + 20:
                left_col.append(block)
            elif x0 > mid - 20:
                right_col.append(block)
            else:
                full_width.append(block)

        # Sort each column by Y position
        full_width.sort(key=lambda b: b["bbox"][1])
        left_col.sort(key=lambda b: b["bbox"][1])
        right_col.sort(key=lambda b: b["bbox"][1])

        # Build interleaved output: full-width blocks, then column pairs
        all_elements = []
        for b in full_width:
            all_elements.append(("full", b["bbox"][1], b))
        if left_col or right_col:
            min_y = min(
                (left_col[0]["bbox"][1] if left_col else 9999),
                (right_col[0]["bbox"][1] if right_col else 9999)
            )
            all_elements.append(("cols", min_y, left_col, right_col))

        all_elements.sort(key=lambda x: x[1])

        for item in all_elements:
            if item[0] == "full":
                self._add_block(doc, item[2], font_stats)
            elif item[0] == "cols":
                self._add_column_table(doc, item[2], item[3], font_stats)

    def _add_column_table(self, doc: Document, left_blocks: list,
                          right_blocks: list, font_stats: dict):
        """Render two columns as a borderless Word table."""
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True

        # Remove all borders
        self._set_table_borders(table, show=False)

        # Fill left cell
        left_cell = table.cell(0, 0)
        left_cell.paragraphs[0].text = ""
        for i, block in enumerate(left_blocks):
            if i == 0:
                para = left_cell.paragraphs[0]
            else:
                para = left_cell.add_paragraph()
            self._fill_paragraph(para, block, font_stats)

        # Fill right cell
        right_cell = table.cell(0, 1)
        right_cell.paragraphs[0].text = ""
        for i, block in enumerate(right_blocks):
            if i == 0:
                para = right_cell.paragraphs[0]
            else:
                para = right_cell.add_paragraph()
            self._fill_paragraph(para, block, font_stats)

    # ── Block Rendering ──────────────────────────────────────────

    def _add_block(self, doc: Document, block: dict, font_stats: dict):
        """Add a text block as a properly styled paragraph."""
        style_type = self._classify_block(block, font_stats)
        text = self._get_block_text(block).strip()

        # Detect bullet/numbered list
        is_bullet = bool(_BULLET_PATTERN.match(text))
        is_numbered = bool(_NUMBER_PATTERN.match(text))

        if style_type == "heading1":
            para = doc.add_heading(level=1)
        elif style_type == "heading2":
            para = doc.add_heading(level=2)
        elif is_bullet:
            para = doc.add_paragraph(style='List Bullet')
        elif is_numbered:
            para = doc.add_paragraph(style='List Number')
        else:
            para = doc.add_paragraph()

        self._fill_paragraph(para, block, font_stats)

    def _classify_block(self, block: dict, font_stats: dict) -> str:
        """Classify a block as heading1, heading2, or body based on font stats."""
        max_size = 0
        is_bold = False

        for line in block.get("lines", []):
            for span in line.get("spans", []):
                max_size = max(max_size, span.get("size", 0))
                if span.get("flags", 0) & 16:
                    is_bold = True

        if max_size >= font_stats["heading1_min"]:
            return "heading1"
        elif max_size >= font_stats["heading2_min"] and is_bold:
            return "heading2"
        elif max_size > font_stats["body_size"] + 1.5 and is_bold:
            return "heading2"
        return "body"

    def _fill_paragraph(self, para, block: dict, font_stats: dict):
        """Fill paragraph with formatted runs preserving font, size, color, bold, italic."""
        # Clear existing runs
        for run in para.runs:
            run.text = ""

        first_run = True
        for line_idx, line in enumerate(block.get("lines", [])):
            # Add line break between lines (except first)
            if line_idx > 0 and not first_run:
                # Check if previous line ends with hyphen (word wrap)
                prev_text = para.runs[-1].text if para.runs else ""
                if prev_text.endswith("-"):
                    # Remove hyphen for word continuation
                    para.runs[-1].text = prev_text[:-1]
                else:
                    # Add space between lines from same block
                    para.add_run(" ")

            for span in line.get("spans", []):
                text = self._sanitize(span.get("text", ""))
                if not text:
                    continue

                first_run = False
                run = para.add_run(text)

                # Font size
                font_size = span.get("size", font_stats["body_size"])
                run.font.size = Pt(font_size)

                # Font color
                color_int = span.get("color", 0)
                r = (color_int >> 16) & 0xFF
                g = (color_int >> 8) & 0xFF
                b = color_int & 0xFF
                run.font.color.rgb = RGBColor(r, g, b)

                # Bold & Italic
                flags = span.get("flags", 0)
                if flags & 16:
                    run.font.bold = True
                if flags & 2:
                    run.font.italic = True

                # Font name (clean up PDF font names)
                font_name = span.get("font", "")
                clean_font = self._clean_font_name(font_name)
                if clean_font:
                    run.font.name = clean_font

        # Paragraph spacing
        pf = para.paragraph_format
        style_type = self._classify_block(block, font_stats)
        if style_type == "heading1":
            pf.space_before = Pt(16)
            pf.space_after = Pt(6)
        elif style_type == "heading2":
            pf.space_before = Pt(10)
            pf.space_after = Pt(4)
        else:
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)
            pf.line_spacing = Pt(14)

    # ── Image Extraction ─────────────────────────────────────────

    def _extract_images(self, doc: Document, page: fitz.Page,
                        pdf_doc: fitz.Document, tmp_dir: str,
                        page_num: int, header_y: float, footer_y: float) -> list:
        """Extract images from the page and add them to the document."""
        img_info_list = page.get_image_info(xrefs=True)
        seen_xrefs = set()
        positions = []

        for info in img_info_list:
            xref = info.get("xref", 0)
            if xref in seen_xrefs or xref == 0:
                continue
            seen_xrefs.add(xref)

            bbox = info.get("bbox", [])
            if not bbox or len(bbox) < 4:
                continue

            w_pt = abs(bbox[2] - bbox[0])
            h_pt = abs(bbox[3] - bbox[1])
            if w_pt < self.MIN_IMAGE_SIZE or h_pt < self.MIN_IMAGE_SIZE:
                continue

            # Skip images in header/footer zones
            img_y = bbox[1]
            if img_y < header_y or img_y >= footer_y:
                continue

            img_path = os.path.join(tmp_dir, f"img_p{page_num}_x{xref}.png")
            try:
                clip = fitz.Rect(bbox) & page.rect
                if clip.is_empty or clip.width < 5 or clip.height < 5:
                    continue
                # High-quality render
                scale = min(3.0, max(2.0, 600.0 / max(clip.width, clip.height)))
                mat = fitz.Matrix(scale, scale)
                pix = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
                pix.save(img_path)
            except Exception as e:
                logger.warning("Image xref=%d render failed: %s", xref, str(e)[:80])
                continue

            if not os.path.isfile(img_path):
                continue

            # Calculate width in inches
            max_w_in = (page.rect.width - 80) / 72.0
            width_in = min(w_pt / 72.0, max_w_in)
            if width_in < 0.3:
                width_in = min(2.0, max_w_in)

            try:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(img_path, width=Inches(width_in))
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after = Pt(4)
                positions.append(img_y)
                logger.info("Added image xref=%d (%.1f in wide)", xref, width_in)
            except Exception as e:
                logger.warning("Image xref=%d insert failed: %s", xref, str(e)[:80])

        return positions

    # ── Table border helpers ─────────────────────────────────────

    def _set_table_borders(self, table, show: bool = True):
        """Set or remove all table borders."""
        tbl = table._tbl
        tbl_pr = tbl.find(f"{{{W_NS}}}tblPr")
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl.insert(0, tbl_pr)

        existing = tbl_pr.find(f"{{{W_NS}}}tblBorders")
        if existing is not None:
            tbl_pr.remove(existing)

        borders = OxmlElement('w:tblBorders')
        for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f'w:{edge}')
            if show:
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), 'D0D0D0')
            else:
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:color'), 'FFFFFF')
            border.set(qn('w:space'), '0')
            borders.append(border)
        tbl_pr.append(borders)

    # ── Utilities ────────────────────────────────────────────────

    @staticmethod
    def _sanitize(text: str) -> str:
        return _CONTROL_CHARS.sub('', text)

    @staticmethod
    def _clean_font_name(font_name: str) -> str:
        """Clean PDF font name to a standard name."""
        if not font_name:
            return ""
        # Remove subset prefix (e.g., ABCDEF+Arial -> Arial)
        if "+" in font_name:
            font_name = font_name.split("+", 1)[-1]
        # Remove style suffix (e.g., Arial-Bold -> Arial)
        base = font_name.split("-")[0].split(",")[0]
        # Map common PDF font names to Windows/system equivalents
        font_map = {
            "TimesNewRoman": "Times New Roman",
            "ArialMT": "Arial",
            "CourierNew": "Courier New",
            "Helvetica": "Arial",
            "HelveticaNeue": "Arial",
            "Times": "Times New Roman",
            "Courier": "Courier New",
        }
        return font_map.get(base, base)

    def _block_has_text(self, block: dict) -> bool:
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                if self._sanitize(span.get("text", "")).strip():
                    return True
        return False

    def _get_block_text(self, block: dict) -> str:
        parts = []
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                t = self._sanitize(span.get("text", ""))
                if t:
                    parts.append(t)
        return " ".join(parts)
