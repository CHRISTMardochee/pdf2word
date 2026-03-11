"""
Combined Converter module (V2).
Creates a DOCX with each PDF page rendered as a behind-text background image,
with editable text boxes positioned at the exact PDF text block coordinates.

This produces output that is:
- Visually identical to the original PDF (background image provides all visuals)
- Fully editable (text in transparent text boxes at exact coordinates)
"""

import logging
import os
import re
import tempfile

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

logger = logging.getLogger(__name__)

# Regex to strip XML-incompatible control characters
_CONTROL_CHARS = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]')

# Conversion factor: 1 PDF point = 12700 EMU
PT_TO_EMU = 12700


class CombinedConverter:
    """
    Creates a DOCX with:
    1. Each PDF page as a behind-text background image (visual fidelity)
    2. Transparent editable text boxes at exact PDF coordinates (editability)
    """

    def __init__(self, dpi: int = 200):
        self.dpi = dpi
        self._shape_id = 0

    def _next_shape_id(self):
        self._shape_id += 1
        return self._shape_id

    def convert(self, input_pdf: str, output_docx: str,
                pages: list[int] | None = None) -> str:
        logger.info("Combined V2 converting: %s -> %s (dpi=%d)",
                     input_pdf, output_docx, self.dpi)

        pdf_doc = fitz.open(input_pdf)
        doc = Document()

        # Remove default empty paragraph
        if doc.paragraphs:
            doc.paragraphs[0]._element.getparent().remove(
                doc.paragraphs[0]._element)

        page_numbers = pages if pages else list(range(len(pdf_doc)))

        with tempfile.TemporaryDirectory() as tmp_dir:
            for page_idx, page_num in enumerate(page_numbers):
                if page_num >= len(pdf_doc):
                    continue

                page = pdf_doc[page_num]
                logger.info("Processing page %d/%d",
                            page_idx + 1, len(page_numbers))

                # Render page as image
                img_path = os.path.join(tmp_dir, f"page_{page_num}.png")
                zoom = self.dpi / 72.0
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                pix.save(img_path)

                # Page dimensions
                page_w_pt = page.rect.width
                page_h_pt = page.rect.height
                page_w_in = page_w_pt / 72.0
                page_h_in = page_h_pt / 72.0
                page_w_emu = int(page_w_pt * PT_TO_EMU)
                page_h_emu = int(page_h_pt * PT_TO_EMU)

                # Section setup
                if page_idx == 0:
                    section = doc.sections[0]
                else:
                    section = doc.add_section()

                section.page_width = Emu(page_w_emu)
                section.page_height = Emu(page_h_emu)
                section.left_margin = Cm(0)
                section.right_margin = Cm(0)
                section.top_margin = Cm(0)
                section.bottom_margin = Cm(0)

                if page_w_in > page_h_in:
                    section.orientation = WD_ORIENT.LANDSCAPE
                else:
                    section.orientation = WD_ORIENT.PORTRAIT

                # Add background image (behind text, full page)
                bg_para = doc.add_paragraph()
                bg_para.paragraph_format.space_before = Pt(0)
                bg_para.paragraph_format.space_after = Pt(0)
                run = bg_para.add_run()
                # Add inline picture first, then we'll have the relationship
                inline_pic = run.add_picture(img_path, width=Emu(page_w_emu))

                # Convert inline picture to behind-text anchor
                self._convert_inline_to_behind_text(
                    run, page_w_emu, page_h_emu)

                # Extract text blocks and add as positioned text boxes
                text_dict = page.get_text("dict")
                blocks = text_dict.get("blocks", [])

                text_para = doc.add_paragraph()
                text_para.paragraph_format.space_before = Pt(0)
                text_para.paragraph_format.space_after = Pt(0)

                for block in blocks:
                    if block["type"] != 0:
                        continue

                    bbox = block["bbox"]  # (x0, y0, x1, y1) in PDF points

                    # Collect spans for this block
                    spans_data = []
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            text = self._sanitize(span.get("text", ""))
                            if text.strip():
                                spans_data.append({
                                    "text": text,
                                    "size": span.get("size", 10),
                                    "color": span.get("color", 0),
                                    "flags": span.get("flags", 0),
                                    "font": span.get("font", ""),
                                })

                    if not spans_data:
                        continue

                    # Create positioned text box for this block
                    x_emu = int(bbox[0] * PT_TO_EMU)
                    y_emu = int(bbox[1] * PT_TO_EMU)
                    w_emu = int((bbox[2] - bbox[0]) * PT_TO_EMU)
                    h_emu = int((bbox[3] - bbox[1]) * PT_TO_EMU)

                    # Add some padding to width to avoid text wrapping issues
                    w_emu = int(w_emu * 1.05)

                    self._add_text_box(text_para, spans_data,
                                       x_emu, y_emu, w_emu, h_emu)

        pdf_doc.close()
        doc.save(output_docx)
        logger.info("Combined V2 conversion complete: %s", output_docx)
        return output_docx

    @staticmethod
    def _sanitize(text: str) -> str:
        return _CONTROL_CHARS.sub('', text)

    def _convert_inline_to_behind_text(self, run, page_w_emu, page_h_emu):
        """
        Convert an inline picture in a run to a behind-text anchored picture
        at position (0, 0) covering the full page.
        """
        # Find the w:drawing element
        drawing = run._element.find(qn('w:drawing'))
        if drawing is None:
            return

        # Find the inline element
        inline = drawing.find(qn('wp:inline'))
        if inline is None:
            return

        # Get the graphic element from inline
        graphic = inline.find(qn('a:graphic'))
        extent = inline.find(qn('wp:extent'))
        doc_pr = inline.find(qn('wp:docPr'))

        if graphic is None or extent is None:
            return

        # Create anchor element
        anchor = OxmlElement('wp:anchor')
        anchor.set('distT', '0')
        anchor.set('distB', '0')
        anchor.set('distL', '0')
        anchor.set('distR', '0')
        anchor.set('simplePos', '0')
        anchor.set('relativeHeight', '0')  # Behind everything
        anchor.set('behindDoc', '1')  # BEHIND text
        anchor.set('locked', '1')
        anchor.set('layoutInCell', '1')
        anchor.set('allowOverlap', '1')

        # Simple position (unused but required)
        simple_pos = OxmlElement('wp:simplePos')
        simple_pos.set('x', '0')
        simple_pos.set('y', '0')
        anchor.append(simple_pos)

        # Horizontal position - relative to page
        pos_h = OxmlElement('wp:positionH')
        pos_h.set('relativeFrom', 'page')
        pos_h_offset = OxmlElement('wp:posOffset')
        pos_h_offset.text = '0'
        pos_h.append(pos_h_offset)
        anchor.append(pos_h)

        # Vertical position - relative to page
        pos_v = OxmlElement('wp:positionV')
        pos_v.set('relativeFrom', 'page')
        pos_v_offset = OxmlElement('wp:posOffset')
        pos_v_offset.text = '0'
        pos_v.append(pos_v_offset)
        anchor.append(pos_v)

        # Extent (size)
        new_extent = OxmlElement('wp:extent')
        new_extent.set('cx', str(page_w_emu))
        new_extent.set('cy', str(page_h_emu))
        anchor.append(new_extent)

        # Effect extent
        effect_extent = OxmlElement('wp:effectExtent')
        effect_extent.set('l', '0')
        effect_extent.set('t', '0')
        effect_extent.set('r', '0')
        effect_extent.set('b', '0')
        anchor.append(effect_extent)

        # Wrap none (no text wrapping)
        wrap_none = OxmlElement('wp:wrapNone')
        anchor.append(wrap_none)

        # Doc properties
        if doc_pr is not None:
            anchor.append(doc_pr)
        else:
            new_doc_pr = OxmlElement('wp:docPr')
            new_doc_pr.set('id', str(self._next_shape_id()))
            new_doc_pr.set('name', 'Background Image')
            anchor.append(new_doc_pr)

        # Move graphic to anchor
        anchor.append(graphic)

        # Replace inline with anchor in the drawing element
        drawing.remove(inline)
        drawing.append(anchor)

    def _add_text_box(self, para, spans_data, x_emu, y_emu, w_emu, h_emu):
        """
        Add a transparent text box at exact coordinates with formatted text.
        Uses DrawingML anchored shape with wps:txbx.
        """
        WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        A = "http://schemas.openxmlformats.org/drawingml/2006/main"
        WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

        shape_id = self._next_shape_id()

        run = para.add_run()
        drawing = OxmlElement('w:drawing')

        # Anchor element
        anchor = etree.SubElement(drawing, f'{{{WP}}}anchor')
        anchor.set('distT', '0')
        anchor.set('distB', '0')
        anchor.set('distL', '0')
        anchor.set('distR', '0')
        anchor.set('simplePos', '0')
        anchor.set('relativeHeight', str(251658240 + shape_id))
        anchor.set('behindDoc', '0')  # In front of background
        anchor.set('locked', '0')
        anchor.set('layoutInCell', '1')
        anchor.set('allowOverlap', '1')

        # Simple pos
        sp = etree.SubElement(anchor, f'{{{WP}}}simplePos')
        sp.set('x', '0')
        sp.set('y', '0')

        # Position H (relative to page)
        ph = etree.SubElement(anchor, f'{{{WP}}}positionH')
        ph.set('relativeFrom', 'page')
        ph_off = etree.SubElement(ph, f'{{{WP}}}posOffset')
        ph_off.text = str(x_emu)

        # Position V (relative to page)
        pv = etree.SubElement(anchor, f'{{{WP}}}positionV')
        pv.set('relativeFrom', 'page')
        pv_off = etree.SubElement(pv, f'{{{WP}}}posOffset')
        pv_off.text = str(y_emu)

        # Extent
        ext = etree.SubElement(anchor, f'{{{WP}}}extent')
        ext.set('cx', str(w_emu))
        ext.set('cy', str(h_emu))

        # Effect extent
        ee = etree.SubElement(anchor, f'{{{WP}}}effectExtent')
        ee.set('l', '0')
        ee.set('t', '0')
        ee.set('r', '0')
        ee.set('b', '0')

        # Wrap none
        etree.SubElement(anchor, f'{{{WP}}}wrapNone')

        # DocPr
        doc_pr = etree.SubElement(anchor, f'{{{WP}}}docPr')
        doc_pr.set('id', str(shape_id))
        doc_pr.set('name', f'TextBox {shape_id}')

        # Graphic
        graphic = etree.SubElement(anchor, f'{{{A}}}graphic')
        graphic_data = etree.SubElement(graphic, f'{{{A}}}graphicData')
        graphic_data.set('uri',
            'http://schemas.microsoft.com/office/word/2010/wordprocessingShape')

        # WordprocessingShape
        wsp = etree.SubElement(graphic_data, f'{{{WPS}}}wsp')

        # cNvSpPr (text box marker)
        cnv = etree.SubElement(wsp, f'{{{WPS}}}cNvSpPr')
        cnv.set('txBox', '1')

        # Shape properties (transparent, no border)
        sp_pr = etree.SubElement(wsp, f'{{{WPS}}}spPr')
        xfrm = etree.SubElement(sp_pr, f'{{{A}}}xfrm')
        off = etree.SubElement(xfrm, f'{{{A}}}off')
        off.set('x', '0')
        off.set('y', '0')
        ext2 = etree.SubElement(xfrm, f'{{{A}}}ext')
        ext2.set('cx', str(w_emu))
        ext2.set('cy', str(h_emu))

        prst = etree.SubElement(sp_pr, f'{{{A}}}prstGeom')
        prst.set('prst', 'rect')
        etree.SubElement(prst, f'{{{A}}}avLst')

        # No fill (transparent background)
        etree.SubElement(sp_pr, f'{{{A}}}noFill')

        # No border line
        ln = etree.SubElement(sp_pr, f'{{{A}}}ln')
        etree.SubElement(ln, f'{{{A}}}noFill')

        # Text box content
        txbx = etree.SubElement(wsp, f'{{{WPS}}}txbx')
        txbx_content = etree.SubElement(txbx, f'{{{W}}}txbxContent')

        # Add paragraph with spans
        w_p = etree.SubElement(txbx_content, f'{{{W}}}p')

        for span_info in spans_data:
            w_r = etree.SubElement(w_p, f'{{{W}}}r')

            # Run properties — text is INVISIBLE (transparent overlay)
            # The background image provides the visual; this text is for
            # search (Ctrl+F) and editing (click on the text box).
            w_rpr = etree.SubElement(w_r, f'{{{W}}}rPr')

            # Font size — keep original size so text box area is clickable
            font_size_hp = int(span_info["size"] * 2)
            sz = etree.SubElement(w_rpr, f'{{{W}}}sz')
            sz.set(f'{{{W}}}val', str(font_size_hp))
            sz_cs = etree.SubElement(w_rpr, f'{{{W}}}szCs')
            sz_cs.set(f'{{{W}}}val', str(font_size_hp))

            # Color — TRANSPARENT (no visible text on screen)
            # Use fully transparent color via w14:textFill if supported,
            # otherwise use white which is invisible on white backgrounds
            # and nearly invisible on images
            w_color = etree.SubElement(w_rpr, f'{{{W}}}color')
            w_color.set(f'{{{W}}}val', 'FFFFFF')

            # Font name (keep for correct text box sizing)
            font_name = span_info["font"]
            clean_font = font_name.split("+")[-1] if "+" in font_name else font_name
            clean_font = clean_font.split("-")[0] if "-" in clean_font else clean_font
            if clean_font:
                w_fonts = etree.SubElement(w_rpr, f'{{{W}}}rFonts')
                w_fonts.set(f'{{{W}}}ascii', clean_font)
                w_fonts.set(f'{{{W}}}hAnsi', clean_font)

            # Text
            w_t = etree.SubElement(w_r, f'{{{W}}}t')
            w_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            w_t.text = span_info["text"]

        # Body properties (tight margins)
        body_pr = etree.SubElement(wsp, f'{{{WPS}}}bodyPr')
        body_pr.set('wrap', 'square')
        body_pr.set('lIns', '0')
        body_pr.set('tIns', '0')
        body_pr.set('rIns', '0')
        body_pr.set('bIns', '0')

        # Append drawing to run
        run._element.append(drawing)
