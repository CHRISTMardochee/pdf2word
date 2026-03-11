"""
Hybrid Converter module.
Renders PDF pages as high-resolution images and creates a DOCX where each page
is the rendered image with invisible text overlay for search/selectability.

This produces output visually identical to the original PDF while maintaining
editable and searchable text.
"""

import logging
import os
import re
import tempfile

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# Regex to strip XML-incompatible control characters
_CONTROL_CHARS = re.compile(
    r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]'
)


class HybridConverter:
    """
    Converts PDF pages to high-resolution images and places them
    in a DOCX document. Text is extracted from the PDF and placed
    as an invisible overlay on top of each page image for
    search/selection capability.
    """

    def __init__(self, dpi: int = 300, text_overlay: bool = True):
        """
        Args:
            dpi: Resolution for rendering PDF pages (default 300 for print quality).
            text_overlay: Whether to add invisible text overlay for search/select.
        """
        self.dpi = dpi
        self.text_overlay = text_overlay

    def convert(self, input_pdf: str, output_docx: str,
                pages: list[int] | None = None) -> str:
        """
        Convert PDF to DOCX using the hybrid (image + text overlay) approach.

        Args:
            input_pdf: Path to the input PDF.
            output_docx: Path for the output DOCX.
            pages: Optional list of 0-indexed page numbers to convert.

        Returns:
            Path to the generated DOCX.
        """
        logger.info("Hybrid converting: %s -> %s (dpi=%d)", input_pdf, output_docx, self.dpi)

        pdf_doc = fitz.open(input_pdf)
        doc = Document()

        # Remove default empty paragraph
        if doc.paragraphs:
            doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

        page_numbers = pages if pages else list(range(len(pdf_doc)))

        with tempfile.TemporaryDirectory() as tmp_dir:
            for page_idx, page_num in enumerate(page_numbers):
                if page_num >= len(pdf_doc):
                    logger.warning("Page %d does not exist, skipping", page_num)
                    continue

                page = pdf_doc[page_num]
                logger.info("Processing page %d/%d", page_idx + 1, len(page_numbers))

                # Render page as high-resolution image
                img_path = os.path.join(tmp_dir, f"page_{page_num}.png")
                self._render_page(page, img_path)

                # Get page dimensions in inches
                page_width_in = page.rect.width / 72.0
                page_height_in = page.rect.height / 72.0

                # Set up the section for this page
                if page_idx == 0:
                    section = doc.sections[0]
                else:
                    section = doc.add_section()

                # Configure page size and minimal margins
                section.page_width = Inches(page_width_in)
                section.page_height = Inches(page_height_in)
                section.left_margin = Cm(0)
                section.right_margin = Cm(0)
                section.top_margin = Cm(0)
                section.bottom_margin = Cm(0)

                # Handle landscape pages
                if page_width_in > page_height_in:
                    section.orientation = WD_ORIENT.LANDSCAPE
                else:
                    section.orientation = WD_ORIENT.PORTRAIT

                # Add the page image
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.0
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                run = para.add_run()
                run.add_picture(img_path, width=Inches(page_width_in))

                # Add invisible text overlay for searchability
                if self.text_overlay:
                    self._add_text_overlay(doc, page, section)

        pdf_doc.close()
        doc.save(output_docx)
        logger.info("Hybrid conversion complete: %s", output_docx)
        return output_docx

    def _render_page(self, page: fitz.Page, output_path: str):
        """Render a PDF page as a high-resolution PNG image."""
        zoom = self.dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        pix.save(output_path)
        logger.debug("Rendered page: %s (%dx%d px)", output_path, pix.width, pix.height)

    @staticmethod
    def _sanitize_text(text: str) -> str:
        """Remove XML-incompatible control characters from text."""
        return _CONTROL_CHARS.sub('', text)

    def _add_text_overlay(self, doc: Document, page: fitz.Page, section):
        """
        Extract text blocks from the PDF page and add them as a hidden
        text layer in the DOCX. The text is styled as very small, white
        (invisible on white/light backgrounds) to maintain searchability
        without affecting the visual appearance.
        """
        text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        blocks = text_dict.get("blocks", [])

        text_content = []
        for block in blocks:
            if block["type"] != 0:  # Skip non-text blocks
                continue
            for line in block.get("lines", []):
                line_text = ""
                for span in line.get("spans", []):
                    line_text += span.get("text", "")
                cleaned = self._sanitize_text(line_text.strip())
                if cleaned:
                    text_content.append(cleaned)

        if not text_content:
            return

        # Add a hidden paragraph with all the text content for this page
        # This enables search/find functionality in the DOCX
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        full_text = self._sanitize_text(" ".join(text_content))
        run = para.add_run(full_text)

        # Make text invisible: very small, white color
        run.font.size = Pt(1)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Arial"

        logger.debug("Added text overlay: %d chars", len(full_text))

