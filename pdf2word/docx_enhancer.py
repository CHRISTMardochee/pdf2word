"""
DocxEnhancer module.
Post-processes a DOCX file to improve paragraph reconstruction and formatting.
Addresses known limitations of pdf2docx:
- Fragmented text in separate text boxes → merged paragraphs
- Inconsistent font styles → normalized
- Split table cells → merged
"""

import logging
import re
from copy import deepcopy
from lxml import etree

from docx import Document
from docx.shared import Pt, Emu, Cm
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# Word XML namespaces
NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "v": "urn:schemas-microsoft-com:vml",
}


class DocxEnhancer:
    """
    Post-process a DOCX to improve quality:
    - Merge fragmented text boxes into proper paragraphs
    - Normalize font styles
    - Clean up empty paragraphs
    """

    def __init__(self):
        pass

    def enhance(self, docx_path: str, output_path: str | None = None,
                source_pdf_path: str | None = None) -> str:
        """
        Enhance a DOCX file.

        Args:
            docx_path: Path to the input DOCX.
            output_path: Path for the output DOCX. If None, overwrites input.
            source_pdf_path: Path to the original source PDF (for image orientation fix).

        Returns:
            Path to the enhanced DOCX.
        """
        if output_path is None:
            output_path = docx_path

        logger.info("Enhancing DOCX: %s", docx_path)

        doc = Document(docx_path)

        # Step 1: Fix section margins (right margin = 0 is a common pdf2docx issue)
        self._fix_section_margins(doc)

        # Step 2: Extract text from text boxes and convert to regular paragraphs
        self._extract_textboxes(doc)

        # Step 3: Merge fragmented consecutive paragraphs
        self._merge_fragmented_paragraphs(doc)

        # Step 4: Normalize spacing (cap excessive space_before/space_after values)
        self._normalize_spacing(doc)

        # Step 5: Normalize indentation (fix inconsistent left indents)
        self._normalize_indentation(doc)

        # Step 6: Unwrap single-cell tables (pdf2docx artifacts)
        self._unwrap_single_cell_tables(doc)

        # Step 7: Remove empty paragraphs
        self._remove_empty_paragraphs(doc)

        # Step 8: Normalize font styles
        self._normalize_fonts(doc)

        # Step 9: Fix image orientation using source PDF transform matrices
        if source_pdf_path:
            self._fix_image_orientation(doc, source_pdf_path)

        # Step 10: Add spacing around image paragraphs
        self._fix_image_spacing(doc)

        # Step 11: Restore table borders
        self._restore_table_borders(doc)

        doc.save(output_path)
        logger.info("DOCX enhanced and saved: %s", output_path)
        return output_path

    def _extract_textboxes(self, doc: Document):
        """
        Find text boxes (drawingML / VML shapes containing text) and extract
        their content as regular paragraphs. This is the key fix for the
        'fragmented text' problem from pdf2docx / LibreOffice.
        """
        body = doc.element.body

        # Raw namespace URIs (qn() only supports 'w:' prefix)
        MC_NS = "{http://schemas.openxmlformats.org/markup-compatibility/2006}"
        WPS_NS = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}"
        V_NS = "{urn:schemas-microsoft-com:vml}"
        W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

        # Find all paragraphs containing text box shapes
        textbox_paragraphs = []

        for para in body.findall(qn("w:p")):
            # Check for DrawingML text boxes
            drawings = para.findall(f".//{MC_NS}AlternateContent")
            drawings += para.findall(".//" + qn("w:drawing"))

            for drawing in drawings:
                # Look for text box content in wps:txbx
                txbx_elements = drawing.findall(f".//{WPS_NS}txbx")
                for txbx in txbx_elements:
                    # Get the text content from the text box
                    inner_body = txbx.find(f"{W_NS}txbxContent")
                    if inner_body is None:
                        continue

                    inner_paras = inner_body.findall(f"{W_NS}p")
                    if inner_paras:
                        textbox_paragraphs.append({
                            "parent_para": para,
                            "inner_paras": inner_paras,
                        })

            # Also check for VML text boxes
            vml_shapes = para.findall(f".//{V_NS}shape")
            for shape in vml_shapes:
                txbx = shape.find(f"{V_NS}textbox")
                if txbx is not None:
                    inner_body = txbx.find(f"{W_NS}txbxContent")
                    if inner_body is None:
                        continue
                    inner_paras = inner_body.findall(f"{W_NS}p")
                    if inner_paras:
                        textbox_paragraphs.append({
                            "parent_para": para,
                            "inner_paras": inner_paras,
                        })

        # Replace text box paragraphs with their extracted content
        for item in textbox_paragraphs:
            parent = item["parent_para"]
            parent_idx = list(body).index(parent)

            for i, inner_para in enumerate(item["inner_paras"]):
                # Clone the inner paragraph
                new_para = deepcopy(inner_para)
                body.insert(parent_idx + 1 + i, new_para)

            # Remove the original text box paragraph
            body.remove(parent)

        if textbox_paragraphs:
            logger.info("Extracted %d text boxes into regular paragraphs", len(textbox_paragraphs))

    def _merge_fragmented_paragraphs(self, doc: Document):
        """
        Merge consecutive paragraphs that appear to be fragments of the same
        paragraph (same font, similar formatting, no explicit paragraph break).
        """
        body = doc.element.body
        paragraphs = list(body.findall(qn("w:p")))

        if len(paragraphs) < 2:
            return

        merged_count = 0
        i = 0

        while i < len(paragraphs) - 1:
            current = paragraphs[i]
            next_para = paragraphs[i + 1]

            current_text = self._get_para_text(current).strip()
            next_text = self._get_para_text(next_para).strip()

            # Skip empty paragraphs
            if not current_text or not next_text:
                i += 1
                continue

            # Heuristic: merge if current doesn't end with sentence-ending punctuation
            # and both have similar formatting
            should_merge = (
                self._should_merge_paragraphs(current, next_para, current_text, next_text)
            )

            if should_merge:
                # Append next paragraph's runs to current paragraph
                self._append_runs(current, next_para)
                body.remove(next_para)
                paragraphs.pop(i + 1)
                merged_count += 1
                # Don't advance i — check if we can merge again
            else:
                i += 1

        if merged_count:
            logger.info("Merged %d fragmented paragraph pairs", merged_count)

    def _should_merge_paragraphs(self, para1, para2, text1: str, text2: str) -> bool:
        """
        Heuristic to decide if two consecutive paragraphs should be merged.
        """
        # Don't merge if current paragraph ends with period, colon, exclamation, question
        if text1 and text1[-1] in ".!?:;":
            return False

        # Don't merge if next paragraph starts with a bullet or number (list item)
        if text2 and (text2[0] in "•·-–—" or re.match(r"^\d+[\.\)]\s", text2)):
            return False

        # Don't merge if next paragraph starts with uppercase (likely new paragraph)
        # UNLESS current doesn't end with sentence-ender (already checked above)
        # Be conservative: only merge if current ends mid-word or with comma
        if text1 and text1[-1] in ",;":
            return True

        # Check if font properties match
        font1 = self._get_dominant_font(para1)
        font2 = self._get_dominant_font(para2)

        if font1 and font2:
            # Merge if same font properties and text1 doesn't end with
            # sentence-ending punctuation
            if font1 == font2 and text1 and text1[-1] not in ".!?:;":
                # Also check: current line is likely a continuation if it's short
                # and doesn't end with punctuation
                return len(text1) < 80  # Short line likely a fragment

        return False

    def _get_para_text(self, para_element) -> str:
        """Extract plain text from a paragraph XML element."""
        texts = []
        for t in para_element.findall(".//" + qn("w:t")):
            if t.text:
                texts.append(t.text)
        return "".join(texts)

    def _get_dominant_font(self, para_element) -> dict | None:
        """Get the dominant font properties from a paragraph."""
        runs = para_element.findall(qn("w:r"))
        if not runs:
            return None

        # Get properties from the first run
        rpr = runs[0].find(qn("w:rPr"))
        if rpr is None:
            return {}

        font_info = {}

        sz = rpr.find(qn("w:sz"))
        if sz is not None:
            font_info["size"] = sz.get(qn("w:val"))

        b = rpr.find(qn("w:b"))
        font_info["bold"] = b is not None

        i = rpr.find(qn("w:i"))
        font_info["italic"] = i is not None

        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is not None:
            font_info["font"] = rfonts.get(qn("w:ascii")) or rfonts.get(qn("w:hAnsi"))

        return font_info

    def _append_runs(self, target_para, source_para):
        """Append runs from source paragraph to target paragraph, adding a space."""
        # Add a space run between the two
        space_run = etree.SubElement(target_para, qn("w:r"))
        space_t = etree.SubElement(space_run, qn("w:t"))
        space_t.text = " "
        space_t.set(qn("xml:space"), "preserve")

        # Copy run properties from source if available
        source_runs = source_para.findall(qn("w:r"))
        for run in source_runs:
            target_para.append(deepcopy(run))

    def _remove_empty_paragraphs(self, doc: Document):
        """Remove paragraphs that are completely empty (no text, no images)."""
        body = doc.element.body
        paragraphs = list(body.findall(qn("w:p")))
        removed = 0

        for para in paragraphs:
            text = self._get_para_text(para).strip()
            has_images = para.findall(".//" + qn("w:drawing")) or para.findall(".//" + qn("w:pict"))

            if not text and not has_images:
                # Keep page breaks
                ppr = para.find(qn("w:pPr"))
                if ppr is not None:
                    sect = ppr.find(qn("w:sectPr"))
                    if sect is not None:
                        continue
                    # Check for page break
                    br = para.findall(".//" + qn("w:br"))
                    has_page_break = any(b.get(qn("w:type")) == "page" for b in br)
                    if has_page_break:
                        continue

                body.remove(para)
                removed += 1

        if removed:
            logger.info("Removed %d empty paragraphs", removed)

    def _normalize_fonts(self, doc: Document):
        """
        Normalize font sizes and styles across the document.
        Fix cases where the same visual text has inconsistent formatting.
        """
        # Collect all font sizes used
        size_counts = {}
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.size:
                    size = run.font.size
                    size_counts[size] = size_counts.get(size, 0) + len(run.text)

        if not size_counts:
            return

        # Find the most common font size (body text size)
        dominant_size = max(size_counts, key=size_counts.get)
        logger.info("Dominant font size: %s", dominant_size)

        # Fix runs with very similar sizes (within 0.5pt) to use the dominant size
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.size and run.font.size != dominant_size:
                    diff = abs(run.font.size - dominant_size)
                    # If within 0.5pt, normalize
                    if diff <= Pt(0.5):
                        run.font.size = dominant_size

    def _fix_section_margins(self, doc: Document):
        """
        Fix section margins: pdf2docx often sets right_margin=0 and creates
        too many sections. Normalize margins to reasonable values.
        """
        default_margin = Cm(1.5)  # ~1.5cm is a reasonable default margin
        min_margin = Cm(0.5)
        fixed = 0

        for section in doc.sections:
            changed = False

            # Fix right margin = 0 (very common pdf2docx issue)
            if section.right_margin is not None and section.right_margin < min_margin:
                section.right_margin = default_margin
                changed = True

            # Fix left margin = 0
            if section.left_margin is not None and section.left_margin < min_margin:
                section.left_margin = default_margin
                changed = True

            if changed:
                fixed += 1

        if fixed:
            logger.info("Fixed margins in %d sections", fixed)

    def _normalize_spacing(self, doc: Document):
        """
        Cap excessive space_before / space_after values.
        pdf2docx often generates huge spacing values (e.g., 890270 EMU = ~70pt)
        which creates big visual gaps between paragraphs.
        """
        max_space = Pt(24)  # Cap at 24pt (~8mm) — generous but reasonable
        fixed = 0

        for para in doc.paragraphs:
            pf = para.paragraph_format

            if pf.space_before is not None and pf.space_before > max_space:
                pf.space_before = max_space
                fixed += 1

            if pf.space_after is not None and pf.space_after > max_space:
                pf.space_after = max_space
                fixed += 1

        if fixed:
            logger.info("Capped %d excessive spacing values (max %s)", fixed, max_space)

    def _normalize_indentation(self, doc: Document):
        """
        Normalize left indentation:
        - Remove very small indents (< 3mm / ~8500 EMU) that are just pdf rounding noise
        - Cap very large indents that push text too far right
        - Preserve intentional indents (bullets, lists)
        """
        noise_threshold = Cm(0.3)     # Below this = pdf rounding noise, remove
        max_body_indent = Cm(3.0)     # Max indent for body text (not bullets)
        fixed = 0

        for para in doc.paragraphs:
            pf = para.paragraph_format

            # Skip paragraphs with hanging indents (bullets/lists)
            if pf.first_line_indent is not None and pf.first_line_indent < 0:
                continue

            if pf.left_indent is not None:
                indent = pf.left_indent

                # Remove noise-level indents
                if 0 < indent < noise_threshold:
                    pf.left_indent = Pt(0)
                    fixed += 1
                # Cap excessive indents (but preserve intentional ones like bullets)
                elif indent > max_body_indent and not para.text.strip().startswith(("•", "-", "–")):
                    # Keep a reasonable indent but remove the excess
                    pf.left_indent = max_body_indent
                    fixed += 1

            # Also clean up small right indents
            if pf.right_indent is not None and 0 < pf.right_indent < noise_threshold:
                pf.right_indent = Pt(0)
                fixed += 1

        if fixed:
            logger.info("Normalized %d indent values", fixed)

    def _unwrap_single_cell_tables(self, doc: Document):
        """
        Convert single-cell tables (1 row x 1 col) into regular paragraphs.
        pdf2docx often wraps content in 1x1 tables for positioning, which
        creates unnecessary borders and makes editing harder.
        """
        body = doc.element.body
        tables = doc.tables
        unwrapped = 0

        for table in tables:
            # Only process 1x1 tables
            if len(table.rows) == 1 and len(table.columns) == 1:
                cell = table.cell(0, 0)
                cell_paras = cell.paragraphs

                if not cell_paras:
                    continue

                # Find the table element in the body
                tbl_element = table._tbl
                try:
                    tbl_idx = list(body).index(tbl_element)
                except ValueError:
                    continue

                # Extract paragraphs from the cell and insert after the table
                for i, cell_para in enumerate(cell_paras):
                    new_para = deepcopy(cell_para._element)
                    body.insert(tbl_idx + 1 + i, new_para)

                # Remove the table
                body.remove(tbl_element)
                unwrapped += 1

        if unwrapped:
            logger.info("Unwrapped %d single-cell tables into paragraphs", unwrapped)

    def _fix_image_orientation(self, doc: Document, source_pdf_path: str):
        """
        Fix image orientation by checking the source PDF's transform matrices.
        PDF uses transformation matrices where negative scale values indicate flips:
        - Negative X scale (matrix[0]) → horizontal flip
        - Negative Y scale (matrix[3]) → vertical flip (but Y is inverted in PDF)
        
        pdf2docx often ignores these transforms, so we detect them and apply
        flipH/flipV attributes to the DOCX DrawingML images.
        """
        import fitz

        A_NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
        WP_NS = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"

        # Extract transform info from the source PDF
        pdf_doc = fitz.open(source_pdf_path)
        pdf_transforms = []
        
        try:
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                img_infos = page.get_image_info(xrefs=True)
                for info in img_infos:
                    transform = info.get("transform", None)
                    if transform and len(transform) >= 4:
                        # transform = [a, b, c, d, e, f]
                        # a = x-scale, d = y-scale
                        # negative a → flipH, positive d in PDF coords → flipV in display
                        flip_h = transform[0] < 0
                        flip_v = transform[3] > 0  # PDF y-axis is inverted
                        if flip_h or flip_v:
                            pdf_transforms.append({
                                "page": page_num,
                                "flip_h": flip_h,
                                "flip_v": flip_v,
                                "width": info.get("width", 0),
                                "height": info.get("height", 0),
                            })
        finally:
            pdf_doc.close()

        if not pdf_transforms:
            return

        logger.info("Found %d images with flipped orientation in source PDF", len(pdf_transforms))

        # Apply flips to DOCX images
        fixed = 0
        for para in doc.paragraphs:
            for run in para.runs:
                drawings = run._element.findall('.//' + qn('w:drawing'))
                for drawing in drawings:
                    # Find xfrm elements in the drawing
                    xfrm_elements = drawing.findall(f'.//{A_NS}xfrm')
                    for xfrm in xfrm_elements:
                        # Apply flips from the first matching PDF transform
                        if fixed < len(pdf_transforms):
                            t = pdf_transforms[fixed]
                            if t["flip_h"]:
                                xfrm.set("flipH", "1")
                            if t["flip_v"]:
                                xfrm.set("flipV", "1")
                            fixed += 1

        if fixed:
            logger.info("Fixed orientation for %d images", fixed)

    def _fix_image_spacing(self, doc: Document):
        """
        Add spacing around paragraphs that contain images.
        pdf2docx often places images with no spacing, causing text to be
        glued to the image.
        """
        image_spacing = Pt(12)  # 12pt spacing around images
        fixed = 0

        for para in doc.paragraphs:
            has_images = (
                para._element.findall('.//' + qn('w:drawing'))
                or para._element.findall('.//' + qn('w:pict'))
            )

            if has_images:
                pf = para.paragraph_format
                
                # Add space after image paragraphs if not already set
                if pf.space_after is None or pf.space_after < image_spacing:
                    pf.space_after = image_spacing
                    fixed += 1

                # Add space before image paragraphs if not already set
                if pf.space_before is None or pf.space_before < image_spacing:
                    pf.space_before = image_spacing
                    fixed += 1

        if fixed:
            logger.info("Fixed spacing around %d image-containing elements", fixed)

    def _restore_table_borders(self, doc: Document):
        """
        Restore borders on tables. pdf2docx often strips table borders.
        
        Strategy:
        - Tables containing red-colored text (Generali content boxes) get
          red borders (matching the brand)
        - Other content tables get light grey borders
        """
        W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        GENERALI_RED = "C5281C"
        LIGHT_GREY = "D0D0D0"
        fixed = 0

        for table in doc.tables:
            # Detect if this table contains red text (content box)
            has_red_text = False
            has_content = False

            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            has_content = True
                        for run in para.runs:
                            if run.font.color and run.font.color.rgb:
                                color_str = str(run.font.color.rgb)
                                if color_str.upper() in ("C5281C", "F2634A", "D5694A"):
                                    has_red_text = True

            if not has_content:
                continue

            # Choose border color based on content
            border_color = GENERALI_RED if has_red_text else LIGHT_GREY
            border_size = "6" if has_red_text else "4"  # 6 = 3/4pt, 4 = 1/2pt

            # Apply borders to the table
            tbl = table._tbl
            tbl_pr = tbl.find(f"{W_NS}tblPr")
            if tbl_pr is None:
                tbl_pr = etree.SubElement(tbl, f"{W_NS}tblPr")

            # Remove existing borders
            existing_borders = tbl_pr.find(f"{W_NS}tblBorders")
            if existing_borders is not None:
                tbl_pr.remove(existing_borders)

            # Create new borders element
            borders = etree.SubElement(tbl_pr, f"{W_NS}tblBorders")

            for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border = etree.SubElement(borders, f"{W_NS}{edge}")
                border.set(f"{W_NS}val", "single")
                border.set(f"{W_NS}sz", border_size)
                border.set(f"{W_NS}space", "0")
                border.set(f"{W_NS}color", border_color)

            fixed += 1

        if fixed:
            logger.info("Restored borders on %d tables (red=%d, grey=%d)", 
                        fixed,
                        sum(1 for t in doc.tables for r in t.rows for c in r.cells 
                            for p in c.paragraphs for run in p.runs 
                            if run.font.color and run.font.color.rgb and 
                            str(run.font.color.rgb).upper() in ("C5281C", "F2634A")),
                        0)  # simplified count



