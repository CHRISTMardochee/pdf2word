"""
Docling ML Converter module.
Uses IBM's Docling library (ML-powered document understanding) to convert
PDFs to DOCX with intelligent structure detection.

Docling provides:
- ML-based layout analysis (titles, paragraphs, lists, tables, images)
- Reading order detection in complex layouts
- Table structure recognition
- OCR for scanned documents

This converter uses Docling for structural analysis and python-docx for
producing high-quality Word output with formatting.
"""

import logging
import os
import tempfile

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


class DoclingConverter:
    """
    PDF-to-DOCX converter powered by IBM Docling ML engine.
    """

    def __init__(self):
        try:
            from docling.document_converter import DocumentConverter
            self._converter_class = DocumentConverter
        except ImportError:
            raise ImportError(
                "Docling is required for the 'docling' mode.\n"
                "Install it with: pip install docling"
            )

    def convert(self, input_pdf: str, output_docx: str,
                pages: list[int] | None = None) -> str:
        """
        Convert PDF to DOCX using Docling ML engine + PyMuPDF visual styling.

        Pipeline:
        1. Docling ML analysis → structured output (headings, tables, lists)
        2. Build Word DOCX from Docling structure
        3. Post-process: apply visual styles from the original PDF using PyMuPDF
           (fonts, colors, backgrounds, images)
        """
        import fitz

        logger.info("Docling ML converting: %s -> %s", input_pdf, output_docx)

        input_pdf_abs = os.path.abspath(input_pdf)
        if not os.path.isfile(input_pdf_abs):
            raise FileNotFoundError(f"Input PDF not found: {input_pdf_abs}")

        # If specific pages requested, extract them first
        actual_pdf = input_pdf_abs
        tmp_pdf = None

        if pages is not None:
            tmp_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
            src = fitz.open(input_pdf_abs)
            dst = fitz.open()
            for p in pages:
                if p < len(src):
                    dst.insert_pdf(src, from_page=p, to_page=p)
            dst.save(tmp_pdf.name)
            dst.close()
            src.close()
            actual_pdf = tmp_pdf.name
            logger.info("Extracted %d pages for Docling processing", len(pages))

        try:
            # Step 1: Run Docling ML analysis
            logger.info("Running Docling ML analysis (this may take a while)...")
            converter = self._converter_class()
            result = converter.convert(actual_pdf)
            docling_doc = result.document

            # Step 2: Build the Word document from Docling's structured output
            doc = self._build_docx(docling_doc)

            # Step 3: Apply visual styles from the original PDF
            logger.info("Applying visual layout styles from source PDF...")
            self._apply_pdf_visual_styles(doc, actual_pdf)

            # Set margins
            for section in doc.sections:
                section.left_margin = Cm(2.0)
                section.right_margin = Cm(2.0)
                section.top_margin = Cm(1.5)
                section.bottom_margin = Cm(1.5)

            doc.save(output_docx)
            logger.info("Docling conversion complete: %s", output_docx)
            return output_docx

        finally:
            if tmp_pdf is not None:
                try:
                    os.unlink(tmp_pdf.name)
                except OSError:
                    pass

    # ── Visual Layout Post-Processing ────────────────────────────

    def _apply_pdf_visual_styles(self, doc: Document, pdf_path: str):
        """
        Post-process the DOCX by extracting visual styles from the original PDF.
        Uses PyMuPDF to read:
        - Font sizes, bold/italic for each paragraph
        - Colored background rectangles (section headers, info boxes)
        - Images
        """
        import fitz

        pdf_doc = fitz.open(pdf_path)

        # Collect ALL visual data from all pages
        visual_data = self._extract_visual_data(pdf_doc)
        pdf_doc.close()

        # Apply fonts and colors to DOCX paragraphs
        self._apply_fonts_to_paragraphs(doc, visual_data)

        # Apply background shading for headings
        self._apply_heading_styles(doc, visual_data)

        logger.info("Visual layout post-processing complete")

    def _extract_visual_data(self, pdf_doc) -> dict:
        """Extract all visual properties from the PDF using PyMuPDF."""
        import fitz

        data = {
            "text_blocks": [],      # All text blocks with font info
            "backgrounds": [],      # Filled rectangles (colored backgrounds)
            "font_stats": {},       # Font size frequency map
        }

        size_counts = {}

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]

            # Extract text blocks with formatting
            text_dict = page.get_text("dict")
            for block in text_dict.get("blocks", []):
                if block.get("type") != 0:
                    continue

                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if not text:
                            continue

                        size = round(span.get("size", 10), 1)
                        flags = span.get("flags", 0)
                        color_int = span.get("color", 0)

                        # Parse flags: bit 0=superscript, 1=italic, 4=bold
                        is_bold = bool(flags & (1 << 4))
                        is_italic = bool(flags & (1 << 1))

                        # Convert color int to RGB hex
                        r = (color_int >> 16) & 0xFF
                        g = (color_int >> 8) & 0xFF
                        b = color_int & 0xFF
                        color_hex = f"{r:02X}{g:02X}{b:02X}"

                        data["text_blocks"].append({
                            "page": page_num,
                            "text": text[:80],
                            "size": size,
                            "bold": is_bold,
                            "italic": is_italic,
                            "color": color_hex,
                            "font": span.get("font", ""),
                            "bbox": span.get("bbox", []),
                        })

                        size_counts[size] = size_counts.get(size, 0) + len(text)

            # Extract filled rectangles (background colors)
            drawings = page.get_drawings()
            for d in drawings:
                fill = d.get("fill")
                rect = d.get("rect")
                if fill is None or rect is None:
                    continue

                dr = fitz.Rect(rect)
                if dr.width < 50 or dr.height < 10:
                    continue

                hex_color = self._rgb_tuple_to_hex(fill)
                if hex_color == "FFFFFF":
                    continue  # Skip white backgrounds

                data["backgrounds"].append({
                    "page": page_num,
                    "rect": rect,
                    "width": dr.width,
                    "height": dr.height,
                    "color": hex_color,
                    "y": dr.y0,
                })

        # Determine dominant body font size
        if size_counts:
            data["font_stats"]["body_size"] = max(size_counts, key=size_counts.get)
        else:
            data["font_stats"]["body_size"] = 10.0

        logger.info("Extracted %d text spans, %d backgrounds from PDF",
                    len(data["text_blocks"]), len(data["backgrounds"]))

        return data

    def _apply_fonts_to_paragraphs(self, doc: Document, visual_data: dict):
        """Match DOCX paragraphs to PDF text blocks and apply fonts."""
        body_size = visual_data["font_stats"].get("body_size", 10.0)

        # Build a lookup of text -> font properties
        text_fonts = {}
        for block in visual_data["text_blocks"]:
            key = block["text"][:40].strip().lower()
            if key and key not in text_fonts:
                text_fonts[key] = block

        applied = 0
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if not para_text:
                continue

            # Try to match paragraph text to PDF text
            key = para_text[:40].strip().lower()
            match = text_fonts.get(key)

            if match:
                for run in para.runs:
                    # Apply font size
                    pdf_size = match["size"]
                    if pdf_size > body_size * 1.3:
                        run.font.size = Pt(pdf_size)
                    else:
                        run.font.size = Pt(body_size)

                    # Apply bold/italic
                    if match["bold"]:
                        run.font.bold = True
                    if match["italic"]:
                        run.font.italic = True

                    # Apply text color (skip black, it's default)
                    color = match["color"]
                    if color != "000000":
                        r = int(color[0:2], 16)
                        g = int(color[2:4], 16)
                        b = int(color[4:6], 16)
                        run.font.color.rgb = RGBColor(r, g, b)

                    run.font.name = "Calibri"
                    applied += 1

        logger.info("Applied PDF fonts to %d runs", applied)

    def _apply_heading_styles(self, doc: Document, visual_data: dict):
        """Apply colored backgrounds to headings that have colored bands behind them."""
        backgrounds = visual_data["backgrounds"]
        if not backgrounds:
            return

        # Identify section header bands (narrow, wide colored rectangles)
        header_bands = [bg for bg in backgrounds
                       if bg["width"] > 200 and 15 < bg["height"] < 50]

        # Identify info boxes (taller colored rectangles)
        info_boxes = [bg for bg in backgrounds
                     if bg["width"] > 200 and bg["height"] > 50]

        styled_count = 0

        for para in doc.paragraphs:
            para_text = para.text.strip()
            if not para_text:
                continue

            # Check if this paragraph matches a header band text
            for block in visual_data["text_blocks"]:
                block_text = block["text"].strip()
                if not block_text:
                    continue

                # Match: paragraph text starts with the block text
                if para_text[:30].lower() == block_text[:30].lower():
                    bbox = block.get("bbox", [])
                    if len(bbox) < 4:
                        continue

                    text_y = bbox[1]
                    text_page = block["page"]

                    # Check if any header band overlaps this text
                    for band in header_bands:
                        if band["page"] != text_page:
                            continue
                        band_rect = band["rect"]
                        # Check Y overlap
                        if (band_rect.y0 - 5) <= text_y <= (band_rect.y1 + 5):
                            # Apply paragraph shading
                            self._set_paragraph_shading(para, band["color"])
                            # Make text white if background is dark
                            if self._is_dark_color(band["color"]):
                                for run in para.runs:
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                                    run.font.bold = True
                            styled_count += 1
                            break

                    # Check info boxes
                    for box in info_boxes:
                        if box["page"] != text_page:
                            continue
                        box_rect = box["rect"]
                        if (box_rect.y0 - 5) <= text_y <= (box_rect.y1 + 5):
                            self._set_paragraph_shading(para, box["color"])
                            styled_count += 1
                            break

                    break  # Only match once per paragraph

        if styled_count:
            logger.info("Applied %d background styles to paragraphs", styled_count)

    @staticmethod
    def _set_paragraph_shading(para, hex_color: str):
        """Apply background shading to a paragraph."""
        pPr = para._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), hex_color)
        shd.set(qn('w:val'), 'clear')
        pPr.append(shd)

    @staticmethod
    def _is_dark_color(hex_color: str) -> bool:
        """Check if a hex color is dark (needs white text)."""
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        luminance = (0.299 * r + 0.587 * g + 0.114 * b) / 255
        return luminance < 0.5

    @staticmethod
    def _rgb_tuple_to_hex(rgb_tuple) -> str:
        """Convert an RGB float tuple (0-1) to a hex string."""
        r = int(min(max(rgb_tuple[0], 0), 1) * 255)
        g = int(min(max(rgb_tuple[1], 0), 1) * 255)
        b = int(min(max(rgb_tuple[2], 0), 1) * 255)
        return f"{r:02X}{g:02X}{b:02X}"

    def _build_docx(self, docling_doc) -> Document:
        """Build a Word document from Docling's structured output."""
        doc = Document()

        # Set default style
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)

        # Get the markdown representation and parse it into structured elements
        md_text = docling_doc.export_to_markdown()
        lines = md_text.split('\n')

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Skip empty lines
            if not stripped:
                i += 1
                continue

            # Image placeholder
            if stripped == '<!-- image -->':
                i += 1
                continue

            # Heading level 1
            if stripped.startswith('# ') and not stripped.startswith('## '):
                text = stripped[2:].strip()
                para = doc.add_heading(text, level=1)
                i += 1
                continue

            # Heading level 2
            if stripped.startswith('## '):
                text = stripped[3:].strip()
                para = doc.add_heading(text, level=2)
                i += 1
                continue

            # Heading level 3
            if stripped.startswith('### '):
                text = stripped[4:].strip()
                para = doc.add_heading(text, level=3)
                i += 1
                continue

            # Table (markdown table format)
            if stripped.startswith('|'):
                # Collect all table lines
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                self._add_markdown_table(doc, table_lines)
                continue

            # List item (bullet)
            if stripped.startswith('- ') or stripped.startswith('* '):
                text = stripped[2:].strip()
                para = doc.add_paragraph(text, style='List Bullet')
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after = Pt(1)
                i += 1
                continue

            # Numbered list
            if len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in '.):':
                text = stripped[2:].strip()
                para = doc.add_paragraph(text, style='List Number')
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after = Pt(1)
                i += 1
                continue

            # Regular paragraph
            para = doc.add_paragraph()
            run = para.add_run(stripped)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            i += 1

        return doc

    def _add_markdown_table(self, doc: Document, table_lines: list):
        """Parse markdown table lines and add as a Word table."""
        # Filter out separator lines (|---|---|)
        data_lines = []
        for line in table_lines:
            # Remove leading/trailing pipe
            cells = [c.strip() for c in line.strip('|').split('|')]
            # Check if it's a separator line (all dashes)
            if all(c.replace('-', '').replace(':', '').strip() == '' for c in cells):
                continue
            # Filter out empty rows
            if any(c for c in cells):
                data_lines.append(cells)

        if not data_lines:
            return

        # Determine dimensions
        max_cols = max(len(row) for row in data_lines)
        row_count = len(data_lines)

        if max_cols == 0 or row_count == 0:
            return

        table = doc.add_table(rows=row_count, cols=max_cols)
        table.autofit = True

        # Set table borders
        self._set_table_borders(table)

        for row_idx, row_data in enumerate(data_lines):
            for col_idx in range(min(len(row_data), max_cols)):
                cell = table.cell(row_idx, col_idx)
                cell_text = row_data[col_idx].strip()

                para = cell.paragraphs[0]
                if cell_text:
                    run = para.add_run(cell_text)
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'

                    # Bold for header row
                    if row_idx == 0:
                        run.font.bold = True

                    # Center bullet marks
                    if cell_text in ('●', '○', '◦', '•'):
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Header row shading
            if row_idx == 0:
                for col_idx in range(min(len(row_data), max_cols)):
                    cell = table.cell(0, col_idx)
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), '008C99')
                    shading.set(qn('w:val'), 'clear')
                    cell._tc.get_or_add_tcPr().append(shading)
                    # White text for header
                    for run in cell.paragraphs[0].runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)

    def _set_table_borders(self, table):
        """Set table borders."""
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        tbl = table._tbl
        tbl_pr = tbl.find(f"{{{W_NS}}}tblPr")
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl.insert(0, tbl_pr)

        existing = tbl_pr.find(f"{{{W_NS}}}tblBorders")
        if existing is not None:
            tbl_pr.remove(existing)

        borders = OxmlElement('w:tblBorders')
        for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), 'D0D0D0')
            border.set(qn('w:space'), '0')
            borders.append(border)
        tbl_pr.append(borders)
